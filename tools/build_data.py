from __future__ import annotations

import argparse
import csv
import hashlib
import json
import math
import os
import re
import shutil
import sys
from collections import Counter, defaultdict
from datetime import datetime, timezone
from pathlib import Path
from typing import Any


PROJECT_ROOT = Path(__file__).resolve().parents[1]
STATIC_ROOT = PROJECT_ROOT / "static"
DATA_DIR = STATIC_ROOT / "data"
ASSET_ROOT = STATIC_ROOT / "assets"
DEFAULT_SOURCE = Path(os.environ.get("CONTENT_SOURCE", r"D:\美股基本面分析"))


try:
    from docx import Document
except Exception:  # pragma: no cover - handled at runtime
    Document = None

try:
    from pypdf import PdfReader
except Exception:  # pragma: no cover - handled at runtime
    PdfReader = None

try:
    import pandas as pd
except Exception:  # pragma: no cover - handled at runtime
    pd = None


TICKER_HINTS = {
    "AMD": ("AMD", "Advanced Micro Devices"),
    "Arista": ("ANET", "Arista Networks"),
    "ANET": ("ANET", "Arista Networks"),
    "Astera": ("ALAB", "Astera Labs"),
    "ALAB": ("ALAB", "Astera Labs"),
    "Qualcomm": ("QCOM", "Qualcomm"),
    "高通": ("QCOM", "Qualcomm"),
    "TSMC": ("TSM", "Taiwan Semiconductor Manufacturing"),
    "台积电": ("TSM", "Taiwan Semiconductor Manufacturing"),
    "Marvell": ("MRVL", "Marvell Technology"),
    "Maevell": ("MRVL", "Marvell Technology"),
    "微软": ("MSFT", "Microsoft"),
    "Microsoft": ("MSFT", "Microsoft"),
    "美光": ("MU", "Micron Technology"),
    "Micron": ("MU", "Micron Technology"),
    "英伟达": ("NVDA", "NVIDIA"),
    "NVIDIA": ("NVDA", "NVIDIA"),
    "Intuitive": ("ISRG", "Intuitive Surgical"),
    "ISRG": ("ISRG", "Intuitive Surgical"),
    "Serve": ("SERV", "Serve Robotics"),
    "SERV": ("SERV", "Serve Robotics"),
    "United": ("UAL", "United Airlines"),
    "联合航空": ("UAL", "United Airlines"),
    "Delta": ("DAL", "Delta Air Lines"),
    "达美": ("DAL", "Delta Air Lines"),
    "CrowdStrike": ("CRWD", "CrowdStrike"),
    "CrowdSrike": ("CRWD", "CrowdStrike"),
    "Crowd": ("CRWD", "CrowdStrike"),
}

THEME_LABELS = {
    "AI agent": "AI Agent",
    "AI PC": "AI PC",
    "网络安全": "网络安全",
    "商业航空": "商业航空",
    "韩国机器人": "机器人",
    "交易": "交易记录",
}

FINANCIAL_KEYWORDS = [
    "股价",
    "市值",
    "收入",
    "营收",
    "增长",
    "毛利",
    "利润",
    "现金流",
    "FCF",
    "NRR",
    "EPS",
    "EBITDA",
    "估值",
    "PE",
    "P/E",
    "PS",
    "P/S",
    "CAGR",
    "TAM",
    "自由现金流",
]

POINT_KEYWORDS = ["核心", "结论", "护城河", "机会", "优势", "驱动", "受益", "弹性", "投资", "定位"]
RISK_KEYWORDS = ["风险", "竞争", "替代", "估值", "限制", "下滑", "放缓", "不确定"]
AI_KEYWORDS = ["AI", "Agent", "GPU", "ASIC", "数据中心", "算力", "自动化", "机器人", "大模型", "推理"]


def clean_text(value: Any) -> str:
    text = "" if value is None else str(value)
    text = re.sub(r"\s+", " ", text.replace("\u3000", " ")).strip()
    return text


def digest(text: str, length: int = 10) -> str:
    return hashlib.sha1(text.encode("utf-8", errors="ignore")).hexdigest()[:length]


def slugify(text: str, fallback: str = "item") -> str:
    raw = text.strip().lower()
    raw = re.sub(r"[^a-z0-9]+", "-", raw)
    raw = raw.strip("-")
    if raw:
        return raw[:80]
    return f"{fallback}-{digest(text, 8)}"


def safe_asset_name(path: Path) -> str:
    stem = slugify(path.stem, "file")
    return f"{stem}-{digest(str(path))}{path.suffix.lower()}"


def rel_parts(path: Path, source_root: Path) -> list[str]:
    try:
        return list(path.relative_to(source_root).parts)
    except ValueError:
        return list(path.parts)


def infer_theme(path: Path, source_root: Path) -> str:
    parts = rel_parts(path, source_root)
    if not parts:
        return "其他"
    first = parts[0]
    return THEME_LABELS.get(first, first)


def infer_ticker_and_name(path: Path, text_blob: str = "") -> tuple[str | None, str]:
    path_probe = str(path)
    for hint, pair in TICKER_HINTS.items():
        if hint.lower() in path_probe.lower():
            return pair
    probe = f"{path} {text_blob[:2000]}"
    for pattern in [
        r"(?:NASDAQ|NYSE)\s*[:：]\s*([A-Z]{1,5})",
        r"\(([A-Z]{1,5})\s*,?\s*(?:NASDAQ|NYSE)",
        r"\b([A-Z]{2,5})\b",
    ]:
        for match in re.finditer(pattern, probe):
            ticker = match.group(1)
            if ticker in {"AI", "PC", "SOC", "CPU", "GPU", "DSP", "ASIC", "FY", "CAGR", "TAM"}:
                continue
            for hint, pair in TICKER_HINTS.items():
                if ticker == pair[0] or hint in probe:
                    return pair
    for hint, pair in TICKER_HINTS.items():
        if hint.lower() in probe.lower():
            return pair
    title = path.stem.replace("_", " ").replace("-", " ")
    title = re.sub(r"\s+", " ", title).strip()
    return None, title


def title_from_paragraphs(path: Path, paragraphs: list[str], fallback_name: str) -> str:
    skip = {"目录", "目 录", "免责声明"}
    for paragraph in paragraphs[:12]:
        compact = paragraph.strip(" -—─")
        if not compact or compact in skip or "仅供" in compact or len(compact) < 2:
            continue
        if len(compact) <= 80:
            return compact
    return fallback_name or path.stem


def is_heading(text: str) -> bool:
    if len(text) > 60:
        return False
    patterns = [
        r"^[一二三四五六七八九十]+[、.]",
        r"^\d+(\.\d+)*[、.\s]",
        r"^第[一二三四五六七八九十\d]+[章节]",
    ]
    return any(re.match(pattern, text) for pattern in patterns)


def extract_sections(paragraphs: list[str]) -> list[dict[str, Any]]:
    sections: list[dict[str, Any]] = []
    current: dict[str, Any] | None = None
    for paragraph in paragraphs:
        if is_heading(paragraph):
            if current:
                current["body"] = clean_text(" ".join(current["body_parts"]))[:900]
                current.pop("body_parts", None)
                sections.append(current)
            current = {"title": paragraph, "body_parts": []}
        elif current:
            current["body_parts"].append(paragraph)
    if current:
        current["body"] = clean_text(" ".join(current["body_parts"]))[:900]
        current.pop("body_parts", None)
        sections.append(current)
    return [s for s in sections if s.get("body")][:18]


def score_text(text: str, keywords: list[str]) -> int:
    upper = text.upper()
    return sum(upper.count(keyword.upper()) for keyword in keywords)


def parse_percent(text: str) -> float | None:
    matches = re.findall(r"(-?\d+(?:\.\d+)?)\s*%", text)
    if not matches:
        return None
    values = [float(v) for v in matches if abs(float(v)) < 500]
    if not values:
        return None
    return values[0]


def clamp(value: float, lo: float = 0, hi: float = 100) -> float:
    return max(lo, min(hi, value))


def row_cells(row: Any) -> list[str]:
    return [clean_text(cell.text) for cell in row.cells]


def read_docx(path: Path) -> dict[str, Any]:
    if Document is None:
        return {"error": "python-docx is not installed"}
    doc = Document(str(path))
    paragraphs = [clean_text(p.text) for p in doc.paragraphs if clean_text(p.text)]
    tables: list[dict[str, Any]] = []
    metrics: list[dict[str, str]] = []

    for index, table in enumerate(doc.tables):
        rows = [row_cells(row) for row in table.rows]
        rows = [[cell for cell in row if cell] for row in rows]
        rows = [row for row in rows if row]
        if not rows:
            continue
        max_cols = max(len(row) for row in rows)
        normalized = [row + [""] * (max_cols - len(row)) for row in rows]
        headers = normalized[0][:8]
        data_rows = [row[:8] for row in normalized[1:26]]
        table_text = " ".join(" ".join(row) for row in normalized[:10])
        is_financial = any(keyword.lower() in table_text.lower() for keyword in FINANCIAL_KEYWORDS)
        tables.append(
            {
                "index": index + 1,
                "headers": headers,
                "rows": data_rows,
                "is_financial": is_financial,
                "row_count": max(0, len(normalized) - 1),
            }
        )

        for row in normalized:
            if len(row) >= 2:
                label = row[0].strip()
                value = row[1].strip()
                if label and value and any(k.lower() in label.lower() for k in FINANCIAL_KEYWORDS):
                    metrics.append({"label": label[:40], "value": value[:120], "source": f"表格 {index + 1}"})
            if len(row) > 2:
                for cell_index, cell in enumerate(row):
                    if any(k.lower() in cell.lower() for k in FINANCIAL_KEYWORDS):
                        metrics.append(
                            {
                                "label": cell[:40],
                                "value": " / ".join(x for x in row[cell_index + 1 : cell_index + 3] if x)[:120],
                                "source": f"表格 {index + 1}",
                            }
                        )

    full_text = "\n".join(paragraphs)
    ticker, name = infer_ticker_and_name(path, full_text)
    title = title_from_paragraphs(path, paragraphs, name)
    sections = extract_sections(paragraphs)
    points = extract_points(paragraphs, POINT_KEYWORDS)
    risks = extract_points(paragraphs, RISK_KEYWORDS)
    scores = build_scores(full_text, metrics, ticker)

    return {
        "ticker": ticker,
        "name": name,
        "title": title,
        "paragraph_count": len(paragraphs),
        "table_count": len(tables),
        "summary": build_summary(paragraphs),
        "sections": sections,
        "key_points": points,
        "risks": risks,
        "metrics": dedupe_metrics(metrics)[:16],
        "tables": tables[:20],
        "scores": scores,
    }


def extract_points(paragraphs: list[str], keywords: list[str]) -> list[str]:
    points: list[str] = []
    for paragraph in paragraphs:
        if len(paragraph) < 18 or len(paragraph) > 240:
            continue
        if any(keyword.lower() in paragraph.lower() for keyword in keywords):
            points.append(paragraph)
        if len(points) >= 8:
            break
    return points


def build_summary(paragraphs: list[str]) -> str:
    candidates = [
        p
        for p in paragraphs
        if 20 <= len(p) <= 260 and "目录" not in p and "免责声明" not in p and not re.match(r"^[一二三四五六七八九十]+[、.]", p)
    ]
    return clean_text(" ".join(candidates[:3]))[:560]


def dedupe_metrics(metrics: list[dict[str, str]]) -> list[dict[str, str]]:
    seen: set[tuple[str, str]] = set()
    output: list[dict[str, str]] = []
    for item in metrics:
        key = (item["label"], item["value"])
        if key in seen:
            continue
        if not item["value"]:
            continue
        seen.add(key)
        output.append(item)
    return output


def build_scores(text: str, metrics: list[dict[str, str]], ticker: str | None) -> dict[str, float]:
    ai_hits = score_text(text, AI_KEYWORDS)
    moat_hits = score_text(text, ["护城河", "龙头", "全球第一", "切换成本", "生态", "平台", "核心", "市占率"])
    risk_hits = score_text(text, RISK_KEYWORDS)
    growth_values: list[float] = []
    margin_values: list[float] = []

    for metric in metrics:
        combined = f"{metric.get('label', '')} {metric.get('value', '')}"
        percent = parse_percent(combined)
        if percent is None:
            continue
        if any(k in combined for k in ["增长", "YoY", "CAGR", "增速"]):
            growth_values.append(percent)
        if any(k in combined for k in ["毛利", "利润率", "现金流", "FCF", "NRR"]):
            margin_values.append(percent)

    growth = clamp((max(growth_values) if growth_values else 15) * 2.2)
    profitability = clamp((max(margin_values) if margin_values else 35) * 1.15)
    ai_exposure = clamp(35 + math.log(ai_hits + 1, 1.7) * 8)
    moat = clamp(30 + math.log(moat_hits + 1, 1.8) * 9)
    risk = clamp(25 + math.log(risk_hits + 1, 1.8) * 8)

    if ticker in {"CRWD", "NVDA", "MRVL", "AMD", "ALAB", "ANET", "TSM", "QCOM"}:
        ai_exposure = clamp(ai_exposure + 12)
    if ticker in {"UAL", "DAL"}:
        ai_exposure = clamp(ai_exposure - 18)
        profitability = clamp(profitability + 5)
    if ticker in {"SERV"}:
        risk = clamp(risk + 15)
    composite = clamp(ai_exposure * 0.24 + growth * 0.24 + profitability * 0.18 + moat * 0.22 + (100 - risk) * 0.12)

    return {
        "综合评分": round(composite, 1),
        "AI暴露度": round(ai_exposure, 1),
        "增长质量": round(growth, 1),
        "盈利质量": round(profitability, 1),
        "护城河": round(moat, 1),
        "风险压力": round(risk, 1),
    }


def copy_asset(path: Path, source_root: Path, copy_files: bool) -> dict[str, Any]:
    extension = path.suffix.lower().lstrip(".")
    asset_type = {
        "pdf": "pdf",
        "docx": "word",
        "mp4": "video",
        "png": "image",
        "jpg": "image",
        "jpeg": "image",
        "csv": "csv",
        "xlsx": "excel",
    }.get(extension, extension or "file")

    theme = infer_theme(path, source_root)
    ticker, name = infer_ticker_and_name(path)
    parts = rel_parts(path, source_root)
    parent = parts[-2] if len(parts) >= 2 else theme
    group_slug = slugify(ticker or parent or theme, "group")
    theme_slug = slugify(theme, "theme")
    target_dir = ASSET_ROOT / theme_slug / group_slug
    target_name = safe_asset_name(path)
    target = target_dir / target_name
    if copy_files:
        target_dir.mkdir(parents=True, exist_ok=True)
        if not target.exists() or target.stat().st_size != path.stat().st_size:
            shutil.copy2(path, target)
    url = f"assets/{theme_slug}/{group_slug}/{target_name}"
    return {
        "id": digest(str(path)),
        "title": path.name,
        "type": asset_type,
        "extension": extension,
        "source_path": str(path),
        "relative_path": "/".join(parts),
        "url": url,
        "size_bytes": path.stat().st_size,
        "theme": theme,
        "ticker": ticker,
        "company_name": name,
    }


def read_pdf_meta(path: Path) -> dict[str, Any]:
    result: dict[str, Any] = {}
    if PdfReader is None:
        return result
    try:
        reader = PdfReader(str(path))
        result["page_count"] = len(reader.pages)
    except Exception as exc:  # pragma: no cover - best effort
        result["pdf_error"] = str(exc)
    return result


def read_tabular_asset(path: Path) -> dict[str, Any]:
    preview: dict[str, Any] = {}
    if pd is None:
        return preview
    try:
        if path.suffix.lower() == ".csv":
            frame = None
            for encoding in ["utf-8-sig", "gbk", "utf-16"]:
                try:
                    frame = pd.read_csv(path, encoding=encoding)
                    break
                except Exception:
                    continue
            if frame is None:
                return preview
            preview["columns"] = [str(c) for c in frame.columns[:16]]
            preview["rows"] = frame.head(20).fillna("").astype(str).values.tolist()
            preview["row_count"] = int(len(frame))
        elif path.suffix.lower() == ".xlsx":
            sheets = pd.read_excel(path, sheet_name=None)
            preview["sheets"] = {}
            for sheet_name, frame in list(sheets.items())[:4]:
                preview["sheets"][str(sheet_name)] = {
                    "columns": [str(c) for c in frame.columns[:16]],
                    "rows": frame.head(20).fillna("").astype(str).values.tolist(),
                    "row_count": int(len(frame)),
                }
    except Exception as exc:  # pragma: no cover - best effort
        preview["table_error"] = str(exc)
    return preview


def group_key_for_asset(asset: dict[str, Any]) -> str:
    ticker = asset.get("ticker")
    if ticker:
        return ticker
    parts = asset["relative_path"].split("/")
    if len(parts) >= 2:
        return f"topic-{slugify(parts[-2], 'topic')}"
    return f"topic-{slugify(asset['title'], 'topic')}"


def company_record_from_group(group_id: str, assets: list[dict[str, Any]], doc_data: dict[str, Any] | None) -> dict[str, Any]:
    first = assets[0]
    ticker = doc_data.get("ticker") if doc_data else first.get("ticker")
    name = doc_data.get("name") if doc_data else first.get("company_name")
    if not name or name == first["title"]:
        name = first.get("company_name") or first["relative_path"].split("/")[-2]
    title = doc_data.get("title") if doc_data else name
    theme_counter = Counter(asset["theme"] for asset in assets)
    theme = theme_counter.most_common(1)[0][0]
    scores = doc_data.get("scores") if doc_data else build_scores(" ".join(a["title"] for a in assets), [], ticker)
    metrics = doc_data.get("metrics", []) if doc_data else []
    sections = doc_data.get("sections", []) if doc_data else []
    summary = doc_data.get("summary", "") if doc_data else ""
    key_points = doc_data.get("key_points", []) if doc_data else []
    risks = doc_data.get("risks", []) if doc_data else []
    tables = doc_data.get("tables", []) if doc_data else []

    return {
        "id": slugify(ticker or group_id, "company"),
        "group_id": group_id,
        "ticker": ticker,
        "name": name,
        "title": title,
        "theme": theme,
        "summary": summary,
        "key_points": key_points,
        "risks": risks,
        "metrics": metrics,
        "sections": sections,
        "tables": tables,
        "scores": scores,
        "assets": sorted(assets, key=lambda item: item["type"]),
    }


def build_catalog(source_root: Path, copy_files: bool = True) -> dict[str, Any]:
    if not source_root.exists():
        raise FileNotFoundError(f"资料文件夹不存在: {source_root}")

    DATA_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_ROOT.mkdir(parents=True, exist_ok=True)

    supported = {".docx", ".pdf", ".mp4", ".png", ".jpg", ".jpeg", ".csv", ".xlsx"}
    files = sorted(
        [
            path
            for path in source_root.rglob("*")
            if path.is_file() and path.suffix.lower() in supported and not path.name.startswith("~$")
        ],
        key=lambda p: str(p),
    )
    assets: list[dict[str, Any]] = []
    doc_payloads: dict[str, dict[str, Any]] = {}
    tabular_assets: list[dict[str, Any]] = []

    for path in files:
        asset = copy_asset(path, source_root, copy_files)
        if path.suffix.lower() == ".pdf":
            asset.update(read_pdf_meta(path))
        elif path.suffix.lower() in {".csv", ".xlsx"}:
            asset["preview"] = read_tabular_asset(path)
            tabular_assets.append(asset)
        elif path.suffix.lower() == ".docx":
            doc_info = read_docx(path)
            asset["document"] = {
                "title": doc_info.get("title"),
                "paragraph_count": doc_info.get("paragraph_count"),
                "table_count": doc_info.get("table_count"),
            }
            asset["ticker"] = doc_info.get("ticker") or asset.get("ticker")
            asset["company_name"] = doc_info.get("name") or asset.get("company_name")
            doc_payloads[asset["id"]] = doc_info
        assets.append(asset)

    grouped: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for asset in assets:
        grouped[group_key_for_asset(asset)].append(asset)

    companies: list[dict[str, Any]] = []
    for group_id, group_assets in grouped.items():
        doc_asset = next((a for a in group_assets if a["type"] == "word" and a["id"] in doc_payloads), None)
        doc_data = doc_payloads.get(doc_asset["id"]) if doc_asset else None
        companies.append(company_record_from_group(group_id, group_assets, doc_data))

    companies.sort(key=lambda item: (item["scores"]["综合评分"], len(item["assets"])), reverse=True)

    themes: list[dict[str, Any]] = []
    by_theme: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for company in companies:
        by_theme[company["theme"]].append(company)
    for theme, items in sorted(by_theme.items()):
        theme_assets = [asset for item in items for asset in item["assets"]]
        overview = next((item for item in items if "总览" in item["title"] or item["ticker"] is None), items[0])
        themes.append(
            {
                "id": slugify(theme, "theme"),
                "name": theme,
                "company_count": sum(1 for item in items if item.get("ticker")),
                "report_count": len(items),
                "asset_count": len(theme_assets),
                "summary": overview.get("summary", ""),
                "companies": [item["id"] for item in items],
                "asset_types": Counter(asset["type"] for asset in theme_assets),
            }
        )

    file_type_counts = Counter(asset["type"] for asset in assets)
    total_size = sum(asset["size_bytes"] for asset in assets)
    catalog = {
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "source_root": str(source_root),
        "stats": {
            "file_count": len(assets),
            "company_count": sum(1 for c in companies if c.get("ticker")),
            "theme_count": len(themes),
            "total_size_bytes": total_size,
            "file_type_counts": dict(file_type_counts),
        },
        "themes": themes,
        "companies": companies,
        "assets": assets,
        "tabular_assets": tabular_assets,
    }
    return catalog


class EnhancedJSONEncoder(json.JSONEncoder):
    def default(self, obj: Any) -> Any:
        if isinstance(obj, Counter):
            return dict(obj)
        return super().default(obj)


def main() -> int:
    parser = argparse.ArgumentParser(description="构建美股基本面分析交互网页数据")
    parser.add_argument("--source", default=str(DEFAULT_SOURCE), help="资料源文件夹")
    parser.add_argument("--output", default=str(DATA_DIR / "catalog.json"), help="输出 JSON 路径")
    parser.add_argument("--no-copy", action="store_true", help="只抽取数据，不复制资产")
    args = parser.parse_args()

    source_root = Path(args.source)
    catalog = build_catalog(source_root, copy_files=not args.no_copy)
    output = Path(args.output)
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(json.dumps(catalog, ensure_ascii=False, indent=2, cls=EnhancedJSONEncoder), encoding="utf-8")

    print(
        json.dumps(
            {
                "ok": True,
                "output": str(output),
                "files": catalog["stats"]["file_count"],
                "companies": catalog["stats"]["company_count"],
                "themes": catalog["stats"]["theme_count"],
                "size_mb": round(catalog["stats"]["total_size_bytes"] / 1024 / 1024, 2),
            },
            ensure_ascii=False,
        )
    )
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as exc:
        print(json.dumps({"ok": False, "error": str(exc)}, ensure_ascii=False), file=sys.stderr)
        raise
